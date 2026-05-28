from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Datos color palette
PURPLE   = RGBColor(0x6D, 0x60, 0xF3)
CTA      = RGBColor(0x6D, 0x2D, 0x90)
DARK     = RGBColor(0x2D, 0x1A, 0x4E)
RED      = RGBColor(0xC0, 0x39, 0x2B)
AMBER    = RGBColor(0xD9, 0x77, 0x06)
GREY     = RGBColor(0x94, 0x92, 0xA8)
TEXT2    = RGBColor(0x4A, 0x47, 0x63)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = PURPLE

def add_body(doc, text, color=TEXT2, size=11, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.italic = italic
    return p

def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2D1A4E')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════
# HEADER
# ════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Datos')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = DARK
r2 = p.add_run(' Health')
r2.font.size = Pt(11); r2.font.color.rgb = GREY

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(4)
r3 = p2.add_run('PriorityPulse')
r3.bold = True; r3.font.size = Pt(28); r3.font.color.rgb = DARK

add_body(doc,
    "A clinician triage layer that turns Datos' Care Team Dashboard from a data viewer "
    "into a prioritized action queue — so care teams know who needs them first.",
    size=12, color=TEXT2)

add_rule(doc)

# ════════════════════════════════════════
# HERO STATS
# ════════════════════════════════════════
doc.add_paragraph()
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'

stats = [
    ('60%+', 'of clinicians report EHR/workflow burden\nas a major driver of burnout', 'Industry research · 2024', 'FEF2F2', RED),
    ('#1',   'Alert fatigue — the top clinician complaint\nacross all Hospital at Home programs', 'AHRQ · JMIR · Springer', 'FFFFFF', DARK),
    ('$15K', 'Average cost of a preventable readmission —\nthe outcome HiTH programs exist to reduce', 'US Healthcare · CMS data', 'FFFFFF', DARK),
]

for i, (num, label, source, bg, num_color) in enumerate(stats):
    cell = tbl.rows[0].cells[i]
    set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p_n = cell.paragraphs[0]
    p_n.paragraph_format.space_before = Pt(10)
    p_n.paragraph_format.space_after = Pt(2)
    p_n.paragraph_format.left_indent = Pt(8)
    r_n = p_n.add_run(num)
    r_n.bold = True; r_n.font.size = Pt(22); r_n.font.color.rgb = num_color

    p_l = cell.add_paragraph(label)
    p_l.paragraph_format.space_after = Pt(4)
    p_l.paragraph_format.left_indent = Pt(8)
    for run in p_l.runs:
        run.font.size = Pt(9.5); run.font.color.rgb = TEXT2

    p_s = cell.add_paragraph(source)
    p_s.paragraph_format.space_after = Pt(10)
    p_s.paragraph_format.left_indent = Pt(8)
    for run in p_s.runs:
        run.font.size = Pt(8); run.font.color.rgb = GREY

# ════════════════════════════════════════
# INSIGHT
# ════════════════════════════════════════
doc.add_paragraph()
insight_tbl = doc.add_table(rows=1, cols=1)
insight_tbl.style = 'Table Grid'
ic = insight_tbl.rows[0].cells[0]
set_cell_bg(ic, 'F7F5FF')

p_i = ic.paragraphs[0]
p_i.paragraph_format.space_before = Pt(8)
p_i.paragraph_format.space_after = Pt(8)
p_i.paragraph_format.left_indent = Pt(10)
p_i.paragraph_format.right_indent = Pt(10)
r1 = p_i.add_run("The Care Team Dashboard still shows ")
r1.font.size = Pt(11); r1.font.color.rgb = DARK
r2 = p_i.add_run("data and alerts, not priorities.")
r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = PURPLE

# ════════════════════════════════════════
# PROBLEM + SOLUTION
# ════════════════════════════════════════
doc.add_paragraph()
ps_tbl = doc.add_table(rows=1, cols=2)
ps_tbl.style = 'Table Grid'

sections_data = [
    ('THE PROBLEM', 'Automated patient workflows, manual clinician triage',
     'Clinicians must review every alert, every patient, every morning — and decide for themselves who matters most. More patients on the platform means more alerts, not more clarity.'),
    ('THE SOLUTION', 'One queue. One question answered.',
     'PriorityPulse adds a triage intelligence layer on top of the existing Care Team Dashboard. It synthesizes vitals trends, and silence patterns into a single composite score — then ranks patients into a prioritized morning queue. One screen. "Who do I call first?" answered in under 30 seconds.'),
]
for i, (lbl, heading, body) in enumerate(sections_data):
    cell = ps_tbl.rows[0].cells[i]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p_lbl = cell.paragraphs[0]
    p_lbl.paragraph_format.space_before = Pt(10)
    p_lbl.paragraph_format.space_after = Pt(3)
    p_lbl.paragraph_format.left_indent = Pt(8)
    r = p_lbl.add_run(lbl)
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = PURPLE

    p_h = cell.add_paragraph(heading)
    p_h.paragraph_format.space_after = Pt(6)
    p_h.paragraph_format.left_indent = Pt(8)
    for run in p_h.runs:
        run.bold = True; run.font.size = Pt(12); run.font.color.rgb = DARK

    p_b = cell.add_paragraph(body)
    p_b.paragraph_format.space_after = Pt(10)
    p_b.paragraph_format.left_indent = Pt(8)
    p_b.paragraph_format.right_indent = Pt(8)
    for run in p_b.runs:
        run.font.size = Pt(10.5); run.font.color.rgb = TEXT2

# ════════════════════════════════════════
# HOW THE LOOP WORKS
# ════════════════════════════════════════
add_label(doc, 'How the loop works')
steps = [
    ('1', 'CareApp prompts daily check-in (personalized to program + recovery day)'),
    ('2', 'Patient submits vitals, taps "okay / concern", or stays silent'),
    ('3', 'Triage engine scores composite risk: vitals + adherence + silence'),
    ('4', 'Care Team Dashboard re-ranks the queue automatically'),
    ('5', 'Care team acts on the right patient first, every morning'),
    ('✓', 'Earlier intervention · fewer readmissions · better outcomes'),
]
for num, text in steps:
    p_step = doc.add_paragraph()
    p_step.paragraph_format.space_after = Pt(3)
    p_step.paragraph_format.left_indent = Pt(8)
    r_num = p_step.add_run(f'  {num}  ')
    r_num.bold = True; r_num.font.size = Pt(10); r_num.font.color.rgb = PURPLE
    r_txt = p_step.add_run(f'  {text}')
    r_txt.font.size = Pt(10.5); r_txt.font.color.rgb = TEXT2

# ════════════════════════════════════════
# TWO SURFACES
# ════════════════════════════════════════
add_label(doc, 'The two surfaces')
feat_tbl = doc.add_table(rows=1, cols=2)
feat_tbl.style = 'Table Grid'

features = [
    ('Patient: Enhanced CareApp Check-In', 'FFFFFF',
     'A small upgrade to the existing CareApp — one personalized daily prompt that references yesterday\'s data. One tap to confirm "okay" or flag "concern." Silence is now a tracked signal, not a void. Built on what already ships. No new app to download, no new login.'),
    ('Care Team: Priority Dashboard', '2D1A4E',
     'A new layer on the existing Care Team Dashboard that answers one question every morning: who needs me first? Patients ranked by composite risk score + silence duration. Silent patients pinned to the top. Click any row for 7-day signal history, latest PROM response, and score breakdown. Under 30 seconds, end-to-end.'),
]
for i, (title, bg, body) in enumerate(features):
    cell = feat_tbl.rows[0].cells[i]
    set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    is_dark = bg == '2D1A4E'

    p_h = cell.paragraphs[0]
    p_h.paragraph_format.space_before = Pt(10)
    p_h.paragraph_format.space_after = Pt(6)
    p_h.paragraph_format.left_indent = Pt(8)
    r = p_h.add_run(title)
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = WHITE if is_dark else DARK

    p_b = cell.add_paragraph(body)
    p_b.paragraph_format.space_after = Pt(10)
    p_b.paragraph_format.left_indent = Pt(8)
    p_b.paragraph_format.right_indent = Pt(8)
    for run in p_b.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xDD) if is_dark else TEXT2

# ════════════════════════════════════════
# MVP SCOPE TABLE
# ════════════════════════════════════════
doc.add_paragraph()
add_label(doc, 'MVP scope — what ships in v1')
add_body(doc,
    "PriorityPulse isn't a rebuild. The CareApp, the data pipeline, and the Care Team Dashboard "
    "already exist. v1 wraps them in a thin triage and prioritization layer that any HiTH program "
    "can turn on through the Design Studio.",
    size=10.5, color=TEXT2, space_after=8)

mvp_tbl = doc.add_table(rows=5, cols=3)
mvp_tbl.style = 'Table Grid'

mvp_headers = ['v1 Capability', 'Already exists?', 'What we add']
for i, h in enumerate(mvp_headers):
    cell = mvp_tbl.rows[0].cells[i]
    set_cell_bg(cell, 'FAFAFE')
    p_h = cell.paragraphs[0]
    p_h.paragraph_format.left_indent = Pt(6)
    r = p_h.add_run(h.upper())
    r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY

mvp_rows = [
    ('Daily CareApp check-in prompt',     'CareApp ships today',                    'One-tap "okay / concern" + silence as a tracked signal'),
    ('Composite triage score',            'Threshold alerts exist; no composite',   'Weighted score: vitals + PROM + adherence + silence'),
    ('Prioritized morning queue',         'Dashboard shows data trends',            'Ranked queue sorted by composite score'),
    ('EMR / device integrations',         'Already in place',                       'No new integration work in v1'),
]
for ri, (cap, exists, add) in enumerate(mvp_rows, start=1):
    row = mvp_tbl.rows[ri]
    for ci, (text, bold, color) in enumerate([
        (cap, True, DARK),
        (exists, False, TEXT2),
        (add, True, PURPLE),
    ]):
        cell = row.cells[ci]
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(10); r.font.color.rgb = color

# ════════════════════════════════════════
# CLINICIAN VOICE QUOTE
# ════════════════════════════════════════
doc.add_paragraph()
q_tbl = doc.add_table(rows=1, cols=1)
q_tbl.style = 'Table Grid'
qc = q_tbl.rows[0].cells[0]
set_cell_bg(qc, 'F9F9FC')

p_q = qc.paragraphs[0]
p_q.paragraph_format.space_before = Pt(10)
p_q.paragraph_format.space_after = Pt(6)
p_q.paragraph_format.left_indent = Pt(12)
p_q.paragraph_format.right_indent = Pt(12)
r_q = p_q.add_run('"I\'m drowning in alerts that don\'t mean anything. I need to know which of my patients actually needs me right now — not see every number that crossed a threshold overnight."')
r_q.italic = True; r_q.font.size = Pt(12); r_q.font.color.rgb = TEXT2

p_attr = qc.add_paragraph('Composite clinician voice — AHRQ, JMIR, Healthcare IT News research')
p_attr.paragraph_format.space_after = Pt(10)
p_attr.paragraph_format.left_indent = Pt(12)
for run in p_attr.runs:
    run.font.size = Pt(8.5); run.bold = True; run.font.color.rgb = GREY

# ════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════
add_rule(doc)
p_link = doc.add_paragraph()
p_link.paragraph_format.space_before = Pt(6)
r_link = p_link.add_run('Live prototype  →  index.html')
r_link.font.size = Pt(10); r_link.font.color.rgb = CTA; r_link.bold = True

out = r'c:\Users\User\Desktop\Ben\Code\Datos\PriorityPulse-One-Pager-v2.docx'
doc.save(out)
print(f'Saved: {out}')

# Convert to PDF
try:
    from docx2pdf import convert
    pdf_out = r'c:\Users\User\Desktop\Ben\Code\Datos\PriorityPulse-One-Pager-v2.pdf'
    convert(out, pdf_out)
    print(f'Saved: {pdf_out}')
except ImportError:
    print('docx2pdf not installed. Run: pip install docx2pdf')
